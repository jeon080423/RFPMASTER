
import io
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def clean_markdown(text):
    """
    Simples markdown cleaner for Word.
    Removes bold/italic markers but keeps the text.
    """
    if not text: return ""
    # Strip <br> tags
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
    # Remove bold/italic ** or *
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    return text

def generate_word_report(results):
    """
    Generates a Word document from the analysis results.
    results: dict { "Section Name": "Content Text", ... }
    """
    doc = Document()
    
    # Set Narrow Margins (0.5 inch / 1.27 cm)
    for section in doc.sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)

    # Title
    title = doc.add_heading('수주비책 (Win Strategy) 분석 보고서', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Set Korean Font Style
    from docx.oxml.ns import qn
    
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    # Add paragraph spacing
    from docx.shared import Pt
    style.paragraph_format.space_after = Pt(6)
    
    for heading in ['Heading 1', 'Heading 2', 'Heading 3']:
        style = doc.styles[heading]
        style.font.name = 'Malgun Gothic'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    
    for section, content in results.items():
        if not content: continue
        
        # Section Header
        doc.add_heading(section, level=1)
        
        lines = content.split('\n')
        table_buffer = []
        in_table = False
        
        for line in lines:
            line = line.strip()
            
            # Detect Table Row
            if line.startswith('|') and line.endswith('|'):
                if not in_table:
                    # Potential start of a table
                    in_table = True
                    table_buffer = [line]
                else:
                    table_buffer.append(line)
                continue
            else:
                # End of a table if we were in one
                if in_table:
                    _process_markdown_table(doc, table_buffer)
                    table_buffer = []
                    in_table = False
            
            # Normal content processing
            if not line:
                continue
                
            if line.startswith('### '):
                doc.add_heading(clean_markdown(line[4:]), level=3)
            elif line.startswith('## '):
                doc.add_heading(clean_markdown(line[3:]), level=2)
            elif line.startswith('# '):
                doc.add_heading(clean_markdown(line[2:]), level=1)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(clean_markdown(line[2:]), style='List Bullet')
            elif line.startswith('1. '):
                doc.add_paragraph(clean_markdown(line[3:]), style='List Number')
            else:
                doc.add_paragraph(clean_markdown(line))
        
        # Flush pending table at end of content
        if in_table and table_buffer:
             _process_markdown_table(doc, table_buffer)
    
    # Final cleanup and save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def _process_markdown_table(doc, lines):
    """
    Parses a markdown table buffer and adds a Word table.
    """
    from docx.oxml.ns import qn
    
    # Filter out divider lines (e.g., |---|---|)
    data_rows = [line for line in lines if not set(line.replace('|', '').strip()) <= set('-: ')]
    
    if not data_rows: return

    # Determine dimensions
    first_row_cells = [c.strip() for c in data_rows[0].strip('|').split('|')]
    cols = len(first_row_cells)
    rows = len(data_rows)
    
    if rows == 0 or cols == 0: return

    # Create Table
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    
    for r, row_text in enumerate(data_rows):
        # Improved row splitting: handle leading/trailing pipes and varying whitespace
        row_content = row_text.strip()
        if row_content.startswith('|'): row_content = row_content[1:]
        if row_content.endswith('|'): row_content = row_content[:-1]
        
        cells = [c.strip() for c in row_content.split('|')]
        
        # Handle mismatch in columns (basic protection)
        for c, text in enumerate(cells):
            if c < cols:
                cell = table.cell(r, c)
                cleaned_text = clean_markdown(text)
                cell.text = cleaned_text
                
                # Clear existing paragraphs in cell and add new ones based on \n
                cell._element.clear_content() # Quick way to clear
                
                # Split by newline and add as paragraphs
                lines_in_cell = cleaned_text.split('\n')
                for i, line_content in enumerate(lines_in_cell):
                    p = cell.add_paragraph(line_content)
                    p.style = doc.styles['Normal']
                    for run in p.runs:
                        run.font.name = 'Malgun Gothic'
                        run.font.size = Pt(9) # Slightly smaller for tables
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
                        if r == 0: # Header Bold
                            run.font.bold = True
